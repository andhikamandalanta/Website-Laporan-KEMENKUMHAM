from flask import Blueprint, render_template, request, send_file
import docx
import io
import os
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Inches

# Blueprint untuk modul Flask
main = Blueprint('main', __name__)

# Fungsi untuk menggantikan kata-kata yang disorot dan gambar dalam dokumen Word
def replace_highlighted_words(doc, replacements, image_paths):
    # Iterasi melalui setiap paragraph dalam dokumen
    for para in doc.paragraphs:
        # Iterasi melalui setiap run (bagian teks yang kontinu) dalam paragraph
        for run in para.runs:
            # Cek apakah teks memiliki highlight berwarna kuning
            if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                # Iterasi melalui setiap kata yang ingin diganti
                for old_word, new_word in replacements.items():
                    # Jika kata lama ditemukan dalam teks run, ganti dengan kata baru
                    if old_word in run.text:
                        run.text = run.text.replace(old_word, new_word)
                        # Hapus highlight warna
                        run.font.highlight_color = None
                # Jika 'gbremonev' ada dalam teks dan path gambarnya tersedia, tambahkan gambar
                if "gbremonev" in run.text and image_paths.get("gbremonev"):
                    run.text = run.text.replace("gbremonev", "")
                    run.add_picture(image_paths["gbremonev"], width=Inches(2))
                    run.font.highlight_color = None
                # Jika 'gmbrsmart' ada dalam teks dan path gambarnya tersedia, tambahkan gambar
                if "gmbrsmart" in run.text and image_paths.get("gmbrsmart"):
                    run.text = run.text.replace("gmbrsmart", "")
                    run.add_picture(image_paths["gmbrsmart"], width=Inches(2))
                    run.font.highlight_color = None
    return doc

# Route untuk halaman utama
@main.route('/')
def home():
    return render_template('form.html')

# Route untuk mengganti kata-kata dan menghasilkan dokumen Word baru
@main.route('/replace', methods=['POST'])
def replace_word():
    # Dictionary untuk menggantikan kata-kata berdasarkan input form
    replacements = {
        'JUDUL': request.form['judul'],
        'Alamat': request.form['alamat'],
        'Telpon': request.form['telpon'],
        'Laman': request.form['laman'],
        'Surel': request.form['surel'],
        'NAMSAT': request.form['namsat'],
        'satker': request.form['satker'],
        'triwulan': request.form['triwulan'],
        'tahun': request.form['tahun'],
        'output': request.form['output'],
        'wulantri': request.form['wulantri'],
        'hunta': request.form['hunta'],
        'putout': request.form['putout'],
        'hambatan1': request.form['hambatan1'],
        'hambatan2': request.form['hambatan2'],
        'rencana1': request.form['rencana1'],
        'rencana2': request.form['rencana2']
    }
    
    # Dictionary untuk menyimpan path dari gambar yang diunggah
    image_paths = {}
    
    # Simpan gambar gbremonev jika diunggah
    if 'gambar_gbremonev' in request.files:
        image_file = request.files['gambar_gbremonev']
        image_path = os.path.join(r'C:\Users\user\Documents\PKL\KEMENKUMHAM\program website\app\static\images\uploads', image_file.filename)
        image_file.save(image_path)
        image_paths['gbremonev'] = image_path

    # Simpan gambar gmbrsmart jika diunggah
    if 'gambar_gmbrsmart' in request.files:
        image_file = request.files['gambar_gmbrsmart']
        image_path = os.path.join(r'C:\Users\user\Documents\PKL\KEMENKUMHAM\program website\app\static\images\uploads', image_file.filename)
        image_file.save(image_path)
        image_paths['gmbrsmart'] = image_path
    
    # Path dari dokumen Word yang akan diupdate
    doc_path = r'C:\Users\user\Documents\PKL\KEMENKUMHAM\program website\file download\laporanevaluasi.docx'
    
    # Jika dokumen tidak ditemukan, beri respons dengan pesan error
    if not os.path.exists(doc_path):
        return f"Dokumen tidak ditemukan di path yang diberikan: {doc_path}"
    
    # Buka dokumen Word yang ada
    doc = docx.Document(doc_path)
    
    # Ganti kata-kata yang disorot dan tambahkan gambar sesuai dengan input yang diberikan
    doc = replace_highlighted_words(doc, replacements, image_paths)
    
    # Path untuk menyimpan dokumen Word yang telah diperbarui
    new_doc_path = r'C:\Users\user\Documents\PKL\KEMENKUMHAM\program website\file download\update_laporanevaluasi.docx'
    
    # Simpan dokumen Word yang telah diperbarui
    doc.save(new_doc_path)
    
    # Kirim dokumen Word yang telah diperbarui sebagai lampiran untuk diunduh
    return send_file(new_doc_path, as_attachment=True)
